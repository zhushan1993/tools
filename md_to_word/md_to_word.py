#!/usr/bin/env python3
"""
Markdown转Word文档转换器（通用版）
基于模板文档，应用正确的样式映射

从P02_ManagementMeasure/script/convert_md_to_docx_final.py移植
重构为接受命令行参数和配置文件
"""

import re
import sys
import json
import argparse
import logging
from pathlib import Path
from docx import Document

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def load_config(config_path=None):
    """
    加载样式映射配置
    """
    default_config = {
        "markdown_to_style": {
            "# ": "文档标题",
            "## ": "一级标题",
            "### ": "二级标题",
            "#### ": "三级标题",
            "普通段落": "缩进正文",
            "表格单元格": "表文字",
            "无序列表": "List Paragraph",
            "有序列表": "List Paragraph"
        },
        "default_styles": {
            "paragraph": "缩进正文",
            "table_cell": "表文字",
            "fallback": "Normal"
        },
        "template_settings": {
            "cover_end_markers": ["一、", "背景", "## ", "# "],
            "preserve_cover": True
        }
    }

    if config_path and Path(config_path).exists():
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                user_config = json.load(f)
                # 合并配置，用户配置优先
                default_config.update(user_config)
                logger.info(f"已加载配置文件: {config_path}")
        except Exception as e:
            logger.warning(f"加载配置文件失败: {e}，使用默认配置")
    else:
        if config_path:
            logger.warning(f"配置文件不存在: {config_path}，使用默认配置")
        else:
            logger.info("使用默认配置")

    return default_config


def markdown_to_docx(md_file_path, template_docx_path, output_docx_path, config=None):
    """
    基于模板转换Markdown为Word文档，应用正确的样式映射

    Args:
        md_file_path: 输入的Markdown文件路径
        template_docx_path: 模板Word文档路径
        output_docx_path: 输出的Word文档路径
        config: 样式映射配置字典
    """
    logger.info(f"开始转换: {md_file_path} -> {output_docx_path}")
    logger.info(f"使用模板: {template_docx_path}")

    if config is None:
        config = load_config()

    # 读取Markdown文件
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except Exception as e:
        logger.error(f"读取Markdown文件失败: {e}")
        return False

    # 打开模板文档
    try:
        doc = Document(template_docx_path)
    except Exception as e:
        logger.error(f"打开模板文档失败: {e}")
        return False

    # 查找封面结束位置（内容开始位置）
    cover_end_index = 0
    cover_end_markers = config.get("template_settings", {}).get("cover_end_markers", ["一、", "背景", "## ", "# "])

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 查找内容开始的标记
        for marker in cover_end_markers:
            if marker in text:
                cover_end_index = i
                break
        if cover_end_index > 0:
            break

    if cover_end_index == 0:
        cover_end_index = len(doc.paragraphs)

    logger.info(f"封面段落: 0-{cover_end_index-1}")
    logger.info(f"内容开始段落: {cover_end_index}")

    # 准备内容段落列表（从内容开始位置到文档末尾）
    content_paragraphs = list(doc.paragraphs)[cover_end_index:]

    # 清空内容段落并重置为默认样式
    default_para_style = config.get("default_styles", {}).get("paragraph", "缩进正文")
    for para in content_paragraphs:
        para.clear()
        try:
            para.style = doc.styles[default_para_style]
        except KeyError:
            try:
                para.style = doc.styles['Normal']
            except KeyError:
                pass

    # 解析Markdown内容
    lines = md_content.split('\n')

    i = 0
    content_para_index = 0  # 当前使用的内容段落索引

    def get_or_add_paragraph(index):
        """获取指定索引的段落，如果不存在则添加"""
        nonlocal content_paragraphs, doc, default_para_style

        if index < len(content_paragraphs):
            return content_paragraphs[index]
        else:
            # 添加新段落
            new_para = doc.add_paragraph()
            content_paragraphs.append(new_para)
            # 新段落默认使用"缩进正文"样式
            try:
                new_para.style = doc.styles[default_para_style]
            except KeyError:
                try:
                    new_para.style = doc.styles['Normal']
                except KeyError:
                    pass
            return new_para

    def apply_style(paragraph, style_name, default_style=None):
        """尝试应用样式，如果不存在则使用默认样式"""
        if default_style is None:
            default_style = config.get("default_styles", {}).get("paragraph", "缩进正文")

        try:
            paragraph.style = doc.styles[style_name]
            return True
        except KeyError:
            try:
                paragraph.style = doc.styles[default_style]
            except KeyError:
                try:
                    paragraph.style = doc.styles['Normal']
                except KeyError:
                    pass
            return False

    def apply_table_cell_style(cell):
        """应用表格单元格样式"""
        table_cell_style = config.get("default_styles", {}).get("table_cell", "表文字")
        # 表格单元格可能包含多个段落，但通常只有一个
        for paragraph in cell.paragraphs:
            try:
                paragraph.style = doc.styles[table_cell_style]
            except KeyError:
                # 如果"表文字"样式不存在，使用Normal
                try:
                    paragraph.style = doc.styles['Normal']
                except KeyError:
                    pass

    logger.info("解析Markdown并应用到文档...")

    while i < len(lines):
        line = lines[i].strip()

        # 空行
        if not line:
            para = get_or_add_paragraph(content_para_index)
            para.clear()
            content_para_index += 1
            i += 1
            continue

        # 文档标题 (#) - 完全跳过，封面已有
        if line.startswith('# '):
            # 跳过文档标题
            i += 1
            continue

        # 一级标题 (##) - 对应样式：一级标题
        if line.startswith('## '):
            title = line[3:].strip()
            # 移除Markdown格式
            title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
            para = get_or_add_paragraph(content_para_index)
            para.text = title
            apply_style(para, '一级标题')
            content_para_index += 1
            i += 1
            continue

        # 二级标题 (###) - 对应样式：二级标题
        if line.startswith('### '):
            title = line[4:].strip()
            title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
            para = get_or_add_paragraph(content_para_index)
            para.text = title
            apply_style(para, '二级标题')
            content_para_index += 1
            i += 1
            continue

        # 三级标题 (####) - 对应样式：三级标题
        if line.startswith('#### '):
            title = line[5:].strip()
            title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
            para = get_or_add_paragraph(content_para_index)
            para.text = title
            apply_style(para, '三级标题')
            content_para_index += 1
            i += 1
            continue

        # 分割线 (---)
        if line == '---':
            para = get_or_add_paragraph(content_para_index)
            para.clear()
            content_para_index += 1
            i += 1
            continue

        # 表格处理
        if '|' in line and line.count('|') >= 2 and not line.startswith('|--'):
            # 收集表格行
            table_lines = []
            j = i

            while j < len(lines) and '|' in lines[j]:
                current_line = lines[j].strip()
                if '|--' in current_line or '|:-' in current_line:
                    j += 1
                    continue
                table_lines.append(current_line)
                j += 1

            if table_lines:
                # 添加一个空段落作为表格前分隔
                para = get_or_add_paragraph(content_para_index)
                para.clear()
                content_para_index += 1

                # 确定列数
                first_row = table_lines[0]
                cols = first_row.count('|') - 1
                rows = len(table_lines)

                # 创建表格
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'

                for row_idx, row_line in enumerate(table_lines):
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < cols:
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            table.cell(row_idx, col_idx).text = cell_text
                            # 应用表格单元格样式
                            apply_table_cell_style(table.cell(row_idx, col_idx))

                i = j
                continue

        # 无序列表
        if line.startswith('- ') or line.startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item = lines[i].strip()[2:].strip()
                item = re.sub(r'\*\*(.*?)\*\*', r'\1', item)
                list_items.append(item)
                i += 1

            for item in list_items:
                para = get_or_add_paragraph(content_para_index)
                para.text = item
                # 列表使用List Paragraph样式
                apply_style(para, 'List Paragraph', default_style='List Paragraph')
                content_para_index += 1

            continue

        # 有序列表
        if re.match(r'^\d+\.\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                item = re.sub(r'^\d+\.\s+', '', lines[i].strip())
                item = re.sub(r'\*\*(.*?)\*\*', r'\1', item)
                list_items.append(item)
                i += 1

            for idx, item in enumerate(list_items):
                para = get_or_add_paragraph(content_para_index)
                para.text = f"{idx+1}. {item}"
                # 有序列表也使用List Paragraph样式
                apply_style(para, 'List Paragraph', default_style='List Paragraph')
                content_para_index += 1

            continue

        # 普通段落：收集连续文本
        paragraph_lines = []
        j = i
        while j < len(lines) and lines[j].strip() and not (
            lines[j].strip().startswith('#') or
            lines[j].strip().startswith('-') or
            lines[j].strip().startswith('*') or
            re.match(r'^\d+\.\s+', lines[j].strip()) or
            ('|' in lines[j] and lines[j].count('|') >= 2)
        ):
            paragraph_lines.append(lines[j].strip())
            j += 1

        if paragraph_lines:
            paragraph_text = ' '.join(paragraph_lines)

            para = get_or_add_paragraph(content_para_index)
            para.clear()

            # 处理加粗文本
            parts = re.split(r'(\*\*.*?\*\*)', paragraph_text)

            for part in parts:
                if not part:
                    continue

                if part.startswith('**') and part.endswith('**'):
                    # 加粗文本
                    bold_text = part[2:-2]
                    run = para.add_run(bold_text)
                    run.bold = True
                else:
                    # 普通文本
                    run = para.add_run(part)

            # 普通段落已默认使用"缩进正文"样式
            content_para_index += 1
            i = j
        else:
            # 如果没有收集到文本，添加空段落
            para = get_or_add_paragraph(content_para_index)
            para.clear()
            content_para_index += 1
            i += 1

    # 保存文档
    try:
        doc.save(output_docx_path)
        logger.info(f"文档已保存: {output_docx_path}")
        return True
    except Exception as e:
        logger.error(f"保存文档失败: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(
        description='将Markdown文件转换为Word文档，基于模板应用样式映射',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s input.md output.docx
  %(prog)s input.md output.docx --template my_template.docx
  %(prog)s input.md output.docx --style-config custom_styles.json --verbose
        """
    )

    parser.add_argument('input', help='输入的Markdown文件路径')
    parser.add_argument('output', help='输出的Word文档路径')
    parser.add_argument('--template', help='模板Word文档路径（可选）')
    parser.add_argument('--style-config', help='样式映射JSON配置文件路径（可选）')
    parser.add_argument('--verbose', action='store_true', help='显示详细输出信息')

    args = parser.parse_args()

    # 设置日志级别
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # 检查输入文件
    input_path = Path(args.input)
    if not input_path.exists():
        logger.error(f"输入文件不存在: {input_path}")
        sys.exit(1)

    # 确定模板文件路径
    if args.template:
        template_path = Path(args.template)
        if not template_path.exists():
            logger.error(f"模板文件不存在: {template_path}")
            sys.exit(1)
    else:
        # 使用默认模板
        script_dir = Path(__file__).parent
        template_path = script_dir / 'config' / 'templates' / 'default.docx'
        if not template_path.exists():
            logger.error(f"默认模板不存在: {template_path}")
            logger.info("请使用 --template 参数指定模板文件")
            sys.exit(1)

    # 加载配置
    config = load_config(args.style_config)

    # 执行转换
    success = markdown_to_docx(
        md_file_path=str(input_path),
        template_docx_path=str(template_path),
        output_docx_path=args.output,
        config=config
    )

    if success:
        logger.info("转换成功！")
        sys.exit(0)
    else:
        logger.error("转换失败！")
        sys.exit(1)


if __name__ == '__main__':
    main()